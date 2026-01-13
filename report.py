from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os

def create_parliamentary_report(
    report_title,
    data_list,
    ai_analysis,
    sector="",
    entity="",
    risk="",
    proposed_action="",
    author="",
    out_dir="reports"
):
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)

    document = Document()

    # Header
    header = document.sections[0].header
    p = header.paragraphs[0]
    p.text = "جمهورية العراق – مجلس النواب\nتقرير رقابي إحصائي وتحليلي"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title = document.add_heading(report_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta = document.add_paragraph()
    meta.add_run(f"القطاع: {sector}\n")
    meta.add_run(f"الجهة: {entity}\n")
    meta.add_run(f"تاريخ التقرير: {datetime.date.today()}\n")
    meta.add_run(f"مُعد التقرير: {author}\n")

    document.add_paragraph("-" * 50)

    # Table
    document.add_heading("أولاً: البيانات الإحصائية", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "المؤشر"
    hdr[1].text = "القيمة"
    hdr[2].text = "الفترة"

    for item in data_list:
        row = table.add_row().cells
        row[0].text = str(item.get("indicator", ""))
        row[1].text = str(item.get("value", ""))
        row[2].text = str(item.get("period", ""))

    # Analysis
    document.add_heading("ثانيًا: التحليل الرقابي", level=2)
    document.add_paragraph(ai_analysis)

    # Risk
    if risk:
        document.add_heading("مستوى الخطورة", level=3)
        document.add_paragraph(risk)

    # Action
    if proposed_action:
        document.add_heading("الإجراء المقترح", level=3)
        document.add_paragraph(proposed_action)

    # Save
    filename = f"{report_title.replace(' ', '_')}_{datetime.date.today()}.docx"
    file_path = os.path.join(out_dir, filename)
    document.save(file_path)

    return file_path
